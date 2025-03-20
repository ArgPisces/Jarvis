import os
import numpy as np
import faiss
from typing import List, Tuple, Optional, Dict
import pickle
from dataclasses import dataclass
from tqdm import tqdm
import fitz  # PyMuPDF for PDF files
from docx import Document as DocxDocument  # python-docx for DOCX files
from pathlib import Path

from yaspin import yaspin
from jarvis.jarvis_platform.registry import PlatformRegistry
import lzma  # 添加 lzma 导入
from threading import Lock
import hashlib

from jarvis.jarvis_utils.config import get_max_paragraph_length, get_max_token_count, get_min_paragraph_length, get_thread_count, get_rag_ignored_paths
from jarvis.jarvis_utils.embedding import get_context_token_count, get_embedding, get_embedding_batch, load_embedding_model
from jarvis.jarvis_utils.output import OutputType, PrettyOutput
from jarvis.jarvis_utils.utils import  get_file_md5, init_env, init_gpu_config

@dataclass
class Document:
    """Document class, for storing document content and metadata"""
    content: str  # Document content
    metadata: Dict  # Metadata (file path, position, etc.)
    md5: str = ""  # File MD5 value, for incremental update detection

class FileProcessor:
    """Base class for file processor"""
    @staticmethod
    def can_handle(file_path: str) -> bool:
        """Determine if the file can be processed"""
        raise NotImplementedError
        
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract file text content"""
        raise NotImplementedError

class TextFileProcessor(FileProcessor):
    """Text file processor"""
    ENCODINGS = ['utf-8', 'gbk', 'gb2312', 'latin1']
    SAMPLE_SIZE = 8192  # Read the first 8KB to detect encoding
    
    @staticmethod
    def can_handle(file_path: str) -> bool:
        """Determine if the file is a text file by trying to decode it"""
        try:
            # Read the first part of the file to detect encoding
            with open(file_path, 'rb') as f:
                sample = f.read(TextFileProcessor.SAMPLE_SIZE)
                
            # Check if it contains null bytes (usually represents a binary file)
            if b'\x00' in sample:
                return False
                
            # Check if it contains too many non-printable characters (usually represents a binary file)
            non_printable = sum(1 for byte in sample if byte < 32 and byte not in (9, 10, 13))  # tab, newline, carriage return
            if non_printable / len(sample) > 0.3:  # If non-printable characters exceed 30%, it is considered a binary file
                return False
                
            # Try to decode with different encodings
            for encoding in TextFileProcessor.ENCODINGS:
                try:
                    sample.decode(encoding)
                    return True
                except UnicodeDecodeError:
                    continue
                    
            return False
            
        except Exception:
            return False
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text content, using the detected correct encoding"""
        detected_encoding = None
        try:
            # First try to detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                
            # Try different encodings
            for encoding in TextFileProcessor.ENCODINGS:
                try:
                    raw_data.decode(encoding)
                    detected_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
                    
            if not detected_encoding:
                raise UnicodeDecodeError(f"Failed to decode file with supported encodings: {file_path}") # type: ignore
                
            # Use the detected encoding to read the file
            with open(file_path, 'r', encoding=detected_encoding, errors='ignore') as f:
                content = f.read()
                
            # Normalize Unicode characters
            import unicodedata
            content = unicodedata.normalize('NFKC', content)
            
            return content
            
        except Exception as e:
            raise Exception(f"Failed to read file: {str(e)}")

class PDFProcessor(FileProcessor):
    """PDF file processor"""
    @staticmethod
    def can_handle(file_path: str) -> bool:
        return Path(file_path).suffix.lower() == '.pdf'
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        text_parts = []
        with fitz.open(file_path) as doc: # type: ignore
            for page in doc:
                text_parts.append(page.get_text()) # type: ignore
        return "\n".join(text_parts)

class DocxProcessor(FileProcessor):
    """DOCX file processor"""
    @staticmethod
    def can_handle(file_path: str) -> bool:
        return Path(file_path).suffix.lower() == '.docx'
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        doc = DocxDocument(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

class RAGTool:
    def __init__(self, root_dir: str):
        """Initialize RAG tool
        
        Args:
            root_dir: Project root directory
        """
        with yaspin(text="初始化环境...", color="cyan") as spinner:
            init_env()
            self.root_dir = root_dir
            os.chdir(self.root_dir)
            spinner.text = "环境初始化完成"
            spinner.ok("✅")
        
        # Initialize configuration
        with yaspin(text="初始化配置...", color="cyan") as spinner:
            self.min_paragraph_length = get_min_paragraph_length()  # Minimum paragraph length
            self.max_paragraph_length = get_max_paragraph_length()  # Maximum paragraph length
            self.context_window = 5  # Fixed context window size
            self.max_token_count = int(get_max_token_count() * 0.8)
            spinner.text = "配置初始化完成"
            spinner.ok("✅")
        
        # Initialize data directory
        with yaspin(text="初始化数据目录...", color="cyan") as spinner:
            self.data_dir = os.path.join(self.root_dir, ".jarvis/rag")
            if not os.path.exists(self.data_dir):
                os.makedirs(self.data_dir)
            spinner.text = "数据目录初始化完成"
            spinner.ok("✅")
            
        # Initialize embedding model
        with yaspin(text="初始化模型...", color="cyan") as spinner:
            try:
                self.embedding_model = load_embedding_model()
                self.vector_dim = self.embedding_model.get_sentence_embedding_dimension()
                spinner.text = "模型加载完成"
                spinner.ok("✅")
            except Exception as e:
                spinner.text = "模型加载失败"
                spinner.fail("❌")
                raise

        with yaspin(text="初始化缓存目录...", color="cyan") as spinner:
            self.cache_dir = os.path.join(self.data_dir, "cache")
            if not os.path.exists(self.cache_dir):
                os.makedirs(self.cache_dir)
                
            self.documents: List[Document] = []
            self.index = None
            self.flat_index = None
            self.file_md5_cache = {}
            spinner.text = "缓存目录初始化完成"
            spinner.ok("✅")
        
        # 加载缓存索引
        self._load_cache_index()

        # Register file processors
        with yaspin(text="初始化文件处理器...", color="cyan") as spinner:
            self.file_processors = [
                TextFileProcessor(),
                PDFProcessor(),
                DocxProcessor()
            ]
            spinner.text = "文件处理器初始化完成"
            spinner.ok("✅")


        # Add thread related configuration
        with yaspin(text="初始化线程配置...", color="cyan") as spinner:
            self.thread_count = get_thread_count()
            self.vector_lock = Lock()  # Protect vector list concurrency
            spinner.text = "线程配置初始化完成"
            spinner.ok("✅")

        # 初始化 GPU 内存配置
        with yaspin(text="初始化 GPU 内存配置...", color="cyan") as spinner:
            with spinner.hidden():
                self.gpu_config = init_gpu_config()
            spinner.text = "GPU 内存配置初始化完成"
            spinner.ok("✅")


    def _get_cache_path(self, file_path: str, cache_type: str = "doc") -> str:
        """Get cache file path for a document
        
        Args:
            file_path: Original file path
            cache_type: Type of cache ("doc" for documents, "vec" for vectors)
            
        Returns:
            str: Cache file path
        """
        # 使用文件路径的哈希作为缓存文件名
        file_hash = hashlib.md5(file_path.encode()).hexdigest()
        
        # 确保不同类型的缓存有不同的目录
        if cache_type == "doc":
            cache_subdir = os.path.join(self.cache_dir, "documents")
        elif cache_type == "vec":
            cache_subdir = os.path.join(self.cache_dir, "vectors")
        else:
            cache_subdir = self.cache_dir
            
        # 确保子目录存在
        if not os.path.exists(cache_subdir):
            os.makedirs(cache_subdir)
            
        return os.path.join(cache_subdir, f"{file_hash}.cache")

    def _load_cache_index(self):
        """Load cache index"""
        index_path = os.path.join(self.data_dir, "index.pkl")
        if os.path.exists(index_path):
            try:
                with yaspin(text="加载缓存索引...", color="cyan") as spinner:
                    with lzma.open(index_path, 'rb') as f:
                        cache_data = pickle.load(f)
                        self.file_md5_cache = cache_data.get("file_md5_cache", {})
                    spinner.text = "缓存索引加载完成"
                    spinner.ok("✅")
                        
                # 从各个缓存文件加载文档
                with yaspin(text="加载缓存文件...", color="cyan") as spinner:
                    for file_path in self.file_md5_cache:
                        doc_cache_path = self._get_cache_path(file_path, "doc")
                        if os.path.exists(doc_cache_path):
                            try:
                                with lzma.open(doc_cache_path, 'rb') as f:
                                    doc_cache_data = pickle.load(f)
                                    self.documents.extend(doc_cache_data["documents"])
                                spinner.text = f"加载文档缓存: {file_path}"
                            except Exception as e:
                                spinner.write(f"❌ 加载文档缓存失败: {file_path}: {str(e)}")
                    spinner.text = "文档缓存加载完成"
                    spinner.ok("✅")
                
                # 重建向量索引
                if self.documents:
                    with yaspin(text="重建向量索引...", color="cyan") as spinner:
                        vectors = []
                        
                        # 按照文档列表顺序加载向量
                        processed_files = set()
                        for doc in self.documents:
                            file_path = doc.metadata['file_path']
                            
                            # 避免重复处理同一个文件
                            if file_path in processed_files:
                                continue
                                
                            processed_files.add(file_path)
                            vec_cache_path = self._get_cache_path(file_path, "vec")
                            
                            if os.path.exists(vec_cache_path):
                                try:
                                    # 加载该文件的向量缓存
                                    with lzma.open(vec_cache_path, 'rb') as f:
                                        vec_cache_data = pickle.load(f)
                                        file_vectors = vec_cache_data["vectors"]
                                    
                                    # 按照文档的chunk_index检索对应向量
                                    doc_indices = [d.metadata['chunk_index'] for d in self.documents 
                                                if d.metadata['file_path'] == file_path]
                                    
                                    # 检查向量数量与文档块数量是否匹配
                                    if len(doc_indices) <= file_vectors.shape[0]:
                                        for idx in doc_indices:
                                            if idx < file_vectors.shape[0]:
                                                vectors.append(file_vectors[idx].reshape(1, -1))
                                    else:
                                        spinner.write(f"⚠️ 向量缓存不匹配: {file_path}")
                                        
                                    spinner.text = f"加载向量缓存: {file_path}"
                                except Exception as e:
                                    spinner.write(f"❌ 加载向量缓存失败: {file_path}: {str(e)}")
                            else:
                                spinner.write(f"⚠️ 缺少向量缓存: {file_path}")
                        
                        if vectors:
                            vectors = np.vstack(vectors)
                            self._build_index(vectors, spinner)
                        spinner.text = f"向量索引重建完成，加载 {len(self.documents)} 个文档片段"
                        spinner.ok("✅")
                                
            except Exception as e:
                PrettyOutput.print(f"加载缓存索引失败: {str(e)}", 
                                output_type=OutputType.WARNING)
                self.documents = []
                self.index = None
                self.flat_index = None
                self.file_md5_cache = {}

    def _save_cache(self, file_path: str, documents: List[Document], vectors: np.ndarray, spinner=None):
        """Save cache for a single file
        
        Args:
            file_path: File path
            documents: List of documents
            vectors: Document vectors
            spinner: Optional spinner for progress display
        """
        try:
            # 保存文档缓存
            if spinner:
                spinner.text = f"保存 {file_path} 的文档缓存..."
            doc_cache_path = self._get_cache_path(file_path, "doc")
            doc_cache_data = {
                "documents": documents
            }
            with lzma.open(doc_cache_path, 'wb') as f:
                pickle.dump(doc_cache_data, f)
                
            # 保存向量缓存
            if spinner:
                spinner.text = f"保存 {file_path} 的向量缓存..."
            vec_cache_path = self._get_cache_path(file_path, "vec")
            vec_cache_data = {
                "vectors": vectors
            }
            with lzma.open(vec_cache_path, 'wb') as f:
                pickle.dump(vec_cache_data, f)
                
            # 更新并保存索引
            if spinner:
                spinner.text = f"更新 {file_path} 的索引缓存..."
            index_path = os.path.join(self.data_dir, "index.pkl")
            index_data = {
                "file_md5_cache": self.file_md5_cache
            }
            with lzma.open(index_path, 'wb') as f:
                pickle.dump(index_data, f)
            
            if spinner:
                spinner.text = f"{file_path} 的缓存保存完成"
                            
        except Exception as e:
            if spinner:
                spinner.text = f"保存 {file_path} 的缓存失败: {str(e)}"
            PrettyOutput.print(f"保存缓存失败: {str(e)}", output_type=OutputType.ERROR)

    def _build_index(self, vectors: np.ndarray, spinner=None):
        """Build FAISS index"""
        if vectors.shape[0] == 0:
            if spinner:
                spinner.text = "向量为空，跳过索引构建"
            self.index = None
            self.flat_index = None
            return
            
        # Create a flat index to store original vectors, for reconstruction
        if spinner:
            spinner.text = "创建平面索引用于向量重建..."
        self.flat_index = faiss.IndexFlatIP(self.vector_dim)
        self.flat_index.add(vectors) # type: ignore
        
        # Create an IVF index for fast search
        if spinner:
            spinner.text = "创建IVF索引用于快速搜索..."
        nlist = max(4, int(vectors.shape[0] / 1000))  # 每1000个向量一个聚类中心
        quantizer = faiss.IndexFlatIP(self.vector_dim)
        self.index = faiss.IndexIVFFlat(quantizer, self.vector_dim, nlist, faiss.METRIC_INNER_PRODUCT)
        
        # Train and add vectors
        if spinner:
            spinner.text = f"训练索引（{vectors.shape[0]}个向量，{nlist}个聚类中心）..."
        self.index.train(vectors) # type: ignore
        
        if spinner:
            spinner.text = "添加向量到索引..."
        self.index.add(vectors) # type: ignore
        
        # Set the number of clusters to probe during search
        if spinner:
            spinner.text = "设置搜索参数..."
        self.index.nprobe = min(nlist, 10)
        
        if spinner:
            spinner.text = f"索引构建完成，共 {vectors.shape[0]} 个向量"

    def _split_text(self, text: str) -> List[str]:
        """Use a more intelligent splitting strategy"""
        # Add overlapping blocks to maintain context consistency
        overlap_size = min(200, self.max_paragraph_length // 4)
        
        paragraphs = []
        current_chunk = []
        current_length = 0
        
        # First split by sentence
        sentences = []
        current_sentence = []
        sentence_ends = {'。', '！', '？', '…', '.', '!', '?'}
        
        for char in text:
            current_sentence.append(char)
            if char in sentence_ends:
                sentence = ''.join(current_sentence)
                if sentence.strip():
                    sentences.append(sentence)
                current_sentence = []
        
        if current_sentence:
            sentence = ''.join(current_sentence)
            if sentence.strip():
                sentences.append(sentence)
        
        # Build overlapping blocks based on sentences
        for sentence in sentences:
            if current_length + len(sentence) > self.max_paragraph_length:
                if current_chunk:
                    chunk_text = ' '.join(current_chunk)
                    if len(chunk_text) >= self.min_paragraph_length:
                        paragraphs.append(chunk_text)
                        
                    # Keep some content as overlap
                    overlap_text = ' '.join(current_chunk[-2:])  # Keep the last two sentences
                    current_chunk = []
                    if overlap_text:
                        current_chunk.append(overlap_text)
                        current_length = len(overlap_text)
                    else:
                        current_length = 0
                        
            current_chunk.append(sentence)
            current_length += len(sentence)
        
        # Process the last chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            if len(chunk_text) >= self.min_paragraph_length:
                paragraphs.append(chunk_text)
        
        return paragraphs


    def _process_file(self, file_path: str, spinner=None) -> List[Document]:
        """Process a single file"""
        try:
            # Calculate file MD5
            if spinner:
                spinner.text = f"计算文件 {file_path} 的MD5..."
            current_md5 = get_file_md5(file_path)
            if not current_md5:
                if spinner:
                    spinner.text = f"文件 {file_path} 计算MD5失败"
                return []

            # Check if the file needs to be reprocessed
            if file_path in self.file_md5_cache and self.file_md5_cache[file_path] == current_md5:
                if spinner:
                    spinner.text = f"文件 {file_path} 未发生变化，跳过处理"
                return []

            # Find the appropriate processor
            if spinner:
                spinner.text = f"查找适用于 {file_path} 的处理器..."
            processor = None
            for p in self.file_processors:
                if p.can_handle(file_path):
                    processor = p
                    break
                    
            if not processor:
                # If no appropriate processor is found, return an empty document
                if spinner:
                    spinner.text = f"没有找到适用于 {file_path} 的处理器，跳过处理"
                return []
            
            # Extract text content
            if spinner:
                spinner.text = f"提取 {file_path} 的文本内容..."
            content = processor.extract_text(file_path)
            if not content.strip():
                if spinner:
                    spinner.text = f"文件 {file_path} 没有文本内容，跳过处理"
                return []
            
            # Split text
            if spinner:
                spinner.text = f"分割 {file_path} 的文本..."
            chunks = self._split_text(content)
            
            # Create document objects
            if spinner:
                spinner.text = f"为 {file_path} 创建 {len(chunks)} 个文档对象..."
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    content=chunk,
                    metadata={
                        "file_path": file_path,
                        "file_type": Path(file_path).suffix.lower(),
                        "chunk_index": i,
                        "total_chunks": len(chunks)
                    },
                    md5=current_md5
                )
                documents.append(doc)
            
            # Update MD5 cache
            self.file_md5_cache[file_path] = current_md5
            if spinner:
                spinner.text = f"文件 {file_path} 处理完成，共创建 {len(documents)} 个文档对象"
            return documents
            
        except Exception as e:
            if spinner:
                spinner.text = f"处理文件失败: {file_path}: {str(e)}"
            PrettyOutput.print(f"处理文件失败: {file_path}: {str(e)}", 
                            output_type=OutputType.ERROR)
            return []

    def _should_ignore_path(self, path: str, ignored_paths: List[str]) -> bool:
        """
        检查路径是否应该被忽略
        
        Args:
            path: 文件或目录路径
            ignored_paths: 忽略模式列表
            
        Returns:
            bool: 如果路径应该被忽略则返回True
        """
        import fnmatch
        import os
        
        # 获取相对路径
        rel_path = path
        if os.path.isabs(path):
            try:
                rel_path = os.path.relpath(path, self.root_dir)
            except ValueError:
                # 如果不能计算相对路径，使用原始路径
                pass
                
        path_parts = rel_path.split(os.sep)
        
        # 检查路径的每一部分是否匹配任意忽略模式
        for part in path_parts:
            for pattern in ignored_paths:
                if fnmatch.fnmatch(part, pattern):
                    return True
                    
        # 检查完整路径是否匹配任意忽略模式
        for pattern in ignored_paths:
            if fnmatch.fnmatch(rel_path, pattern):
                return True
                
        return False
        
    def _is_git_repo(self) -> bool:
        """
        检查当前目录是否为Git仓库
        
        Returns:
            bool: 如果是Git仓库则返回True
        """
        import subprocess
        
        try:
            result = subprocess.run(
                ["git", "rev-parse", "--is-inside-work-tree"],
                cwd=self.root_dir,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                check=False
            )
            return result.returncode == 0 and result.stdout.strip() == "true"
        except Exception:
            return False
    
    def _get_git_managed_files(self) -> List[str]:
        """
        获取Git仓库中被管理的文件列表
        
        Returns:
            List[str]: 被Git管理的文件路径列表（相对路径）
        """
        import subprocess
        
        try:
            # 获取git索引中的文件
            result = subprocess.run(
                ["git", "ls-files"],
                cwd=self.root_dir,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                check=False
            )
            
            if result.returncode != 0:
                return []
                
            git_files = [line.strip() for line in result.stdout.splitlines() if line.strip()]
            
            # 添加未暂存但已跟踪的修改文件
            result = subprocess.run(
                ["git", "ls-files", "--modified"],
                cwd=self.root_dir,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                check=False
            )
            
            if result.returncode == 0:
                modified_files = [line.strip() for line in result.stdout.splitlines() if line.strip()]
                git_files.extend([f for f in modified_files if f not in git_files])
            
            # 转换为绝对路径
            return [os.path.join(self.root_dir, file) for file in git_files]
            
        except Exception as e:
            PrettyOutput.print(f"获取Git管理的文件失败: {str(e)}", output_type=OutputType.WARNING)
            return []

    def build_index(self, dir: str):
        try:
            """Build document index with optimized processing"""
            # Get all files
            with yaspin(text="获取所有文件...", color="cyan") as spinner:
                all_files = []
                
                # 获取需要忽略的路径列表
                ignored_paths = get_rag_ignored_paths()
                
                # 检查是否为Git仓库
                is_git_repo = self._is_git_repo()
                if is_git_repo:
                    git_files = self._get_git_managed_files()
                    # 过滤掉被忽略的文件
                    for file_path in git_files:
                        if self._should_ignore_path(file_path, ignored_paths):
                            continue
                            
                        if os.path.getsize(file_path) > 100 * 1024 * 1024:  # 100MB
                            PrettyOutput.print(f"跳过大文件: {file_path}", 
                                            output_type=OutputType.WARNING)
                            continue
                        all_files.append(file_path)
                else:
                    # 非Git仓库，使用常规文件遍历
                    for root, _, files in os.walk(dir):
                        # 检查目录是否匹配忽略模式
                        if self._should_ignore_path(root, ignored_paths):
                            continue
                            
                        for file in files:
                            file_path = os.path.join(root, file)
                            
                            # 检查文件是否匹配忽略模式
                            if self._should_ignore_path(file_path, ignored_paths):
                                continue
                                
                            if os.path.getsize(file_path) > 100 * 1024 * 1024:  # 100MB
                                PrettyOutput.print(f"跳过大文件: {file_path}", 
                                                output_type=OutputType.WARNING)
                                continue
                            all_files.append(file_path)
                            
                spinner.text = f"获取所有文件完成，共 {len(all_files)} 个文件"
                spinner.ok("✅")

            # Clean up cache for deleted files
            with yaspin(text="清理缓存...", color="cyan") as spinner:
                deleted_files = set(self.file_md5_cache.keys()) - set(all_files)
                deleted_count = len(deleted_files)
                
                if deleted_count > 0:
                    spinner.write(f"🗑️ 删除不存在文件的缓存: {deleted_count} 个")
                    
                for file_path in deleted_files:
                    # Remove from MD5 cache
                    del self.file_md5_cache[file_path]
                    # Remove related documents
                    self.documents = [doc for doc in self.documents if doc.metadata['file_path'] != file_path]
                    # Delete cache files
                    self._delete_file_cache(file_path, None)  # Pass None as spinner to not show individual deletions
                    
                spinner.text = f"清理缓存完成，共删除 {deleted_count} 个不存在文件的缓存"
                spinner.ok("✅")

            # Check file changes
            with yaspin(text="检查文件变化...", color="cyan") as spinner:
                files_to_process = []
                unchanged_files = []
                new_files_count = 0
                modified_files_count = 0
                
                for file_path in all_files:
                    current_md5 = get_file_md5(file_path)
                    if current_md5:  # Only process files that can successfully calculate MD5
                        if file_path in self.file_md5_cache and self.file_md5_cache[file_path] == current_md5:
                            # File未变化，记录但不重新处理
                            unchanged_files.append(file_path)
                        else:
                            # New file or modified file
                            files_to_process.append(file_path)
                            
                            # 如果是修改的文件，删除旧缓存
                            if file_path in self.file_md5_cache:
                                modified_files_count += 1
                                # 删除旧缓存
                                self._delete_file_cache(file_path, spinner)
                                # 从文档列表中移除
                                self.documents = [doc for doc in self.documents if doc.metadata['file_path'] != file_path]
                            else:
                                new_files_count += 1
                
                # 输出汇总信息
                if unchanged_files:
                    spinner.write(f"📚 已缓存文件: {len(unchanged_files)} 个")
                if new_files_count > 0:
                    spinner.write(f"🆕 新增文件: {new_files_count} 个")
                if modified_files_count > 0:
                    spinner.write(f"📝 修改文件: {modified_files_count} 个")
                    
                spinner.text = f"检查文件变化完成，共 {len(files_to_process)} 个文件需要处理"
                spinner.ok("✅")

            # Keep documents for unchanged files
            unchanged_documents = [doc for doc in self.documents 
                                if doc.metadata['file_path'] in unchanged_files]

            # Process files one by one with optimized vectorization
            if files_to_process:
                new_documents = []
                new_vectors = []
                success_count = 0
                skipped_count = 0
                failed_count = 0
                
                with yaspin(text=f"处理文件中 (0/{len(files_to_process)})...", color="cyan") as spinner:
                    for index, file_path in enumerate(files_to_process):
                        spinner.text = f"处理文件中 ({index+1}/{len(files_to_process)}): {file_path}"
                        try:
                            # Process single file
                            file_docs = self._process_file(file_path, spinner)
                            if file_docs:
                                # Vectorize documents from this file
                                spinner.text = f"处理文件中 ({index+1}/{len(files_to_process)}): 为 {file_path} 生成向量嵌入..."
                                texts_to_vectorize = [
                                    f"File:{doc.metadata['file_path']} Content:{doc.content}"
                                    for doc in file_docs
                                ]
                                
                                file_vectors = get_embedding_batch(self.embedding_model, f"({index+1}/{len(files_to_process)}){file_path}", texts_to_vectorize, spinner)
                                
                                # Save cache for this file
                                spinner.text = f"处理文件中 ({index+1}/{len(files_to_process)}): 保存 {file_path} 的缓存..."
                                self._save_cache(file_path, file_docs, file_vectors, spinner)
                                
                                # Accumulate documents and vectors
                                new_documents.extend(file_docs)
                                new_vectors.append(file_vectors)
                                success_count += 1
                            else:
                                # 文件跳过处理
                                skipped_count += 1
                                
                        except Exception as e:
                            spinner.write(f"❌ 处理失败: {file_path}: {str(e)}")
                            failed_count += 1
                    
                    # 输出处理统计
                    spinner.text = f"文件处理完成: 成功 {success_count} 个, 跳过 {skipped_count} 个, 失败 {failed_count} 个"
                    spinner.ok("✅")
                    
                # Update documents list
                self.documents.extend(new_documents)

                # Build final index
                if new_vectors:
                    with yaspin(text="构建最终索引...", color="cyan") as spinner:
                        spinner.text = "合并新向量..."
                        all_new_vectors = np.vstack(new_vectors)
                        
                        unchanged_vector_count = 0
                        if self.flat_index is not None:
                            # Get vectors for unchanged documents
                            spinner.text = "获取未变化文档的向量..."
                            unchanged_vectors = self._get_unchanged_vectors(unchanged_documents, spinner)
                            if unchanged_vectors is not None:
                                unchanged_vector_count = unchanged_vectors.shape[0]
                                spinner.text = f"合并新旧向量（新：{all_new_vectors.shape[0]}，旧：{unchanged_vector_count}）..."
                                final_vectors = np.vstack([unchanged_vectors, all_new_vectors])
                            else:
                                spinner.text = f"仅使用新向量（{all_new_vectors.shape[0]}）..."
                                final_vectors = all_new_vectors
                        else:
                            spinner.text = f"仅使用新向量（{all_new_vectors.shape[0]}）..."
                            final_vectors = all_new_vectors

                        # Build index
                        spinner.text = f"构建索引（向量数量：{final_vectors.shape[0]}）..."
                        self._build_index(final_vectors, spinner)
                        spinner.text = f"索引构建完成，共 {len(self.documents)} 个文档片段"
                        spinner.ok("✅")

                # 输出最终统计信息
                PrettyOutput.print(
                    f"📊 索引统计:\n"
                    f"  • 总文档数: {len(self.documents)} 个文档片段\n"
                    f"  • 已缓存文件: {len(unchanged_files)} 个\n"
                    f"  • 处理文件: {len(files_to_process)} 个\n"
                    f"    - 成功: {success_count} 个\n"
                    f"    - 跳过: {skipped_count} 个\n"
                    f"    - 失败: {failed_count} 个", 
                    OutputType.SUCCESS
                )
        except Exception as e:
            PrettyOutput.print(f"索引构建失败: {str(e)}", 
                            output_type=OutputType.ERROR)

    def _get_unchanged_vectors(self, unchanged_documents: List[Document], spinner=None) -> Optional[np.ndarray]:
        """Get vectors for unchanged documents from existing index"""
        try:
            if not unchanged_documents:
                if spinner:
                    spinner.text = "没有未变化的文档"
                return None

            if spinner:
                spinner.text = f"加载 {len(unchanged_documents)} 个未变化文档的向量..."
            
            # 按文件分组处理
            unchanged_files = set(doc.metadata['file_path'] for doc in unchanged_documents)
            unchanged_vectors = []
            
            for file_path in unchanged_files:
                if spinner:
                    spinner.text = f"加载 {file_path} 的向量..."
                
                # 获取该文件所有文档的chunk索引
                doc_indices = [(i, doc.metadata['chunk_index']) 
                              for i, doc in enumerate(unchanged_documents) 
                              if doc.metadata['file_path'] == file_path]
                
                if not doc_indices:
                    continue
                
                # 加载该文件的向量
                vec_cache_path = self._get_cache_path(file_path, "vec")
                if os.path.exists(vec_cache_path):
                    try:
                        with lzma.open(vec_cache_path, 'rb') as f:
                            vec_cache_data = pickle.load(f)
                            file_vectors = vec_cache_data["vectors"]
                        
                        # 按照chunk_index加载对应的向量
                        for _, chunk_idx in doc_indices:
                            if chunk_idx < file_vectors.shape[0]:
                                unchanged_vectors.append(file_vectors[chunk_idx].reshape(1, -1))
                            
                        if spinner:
                            spinner.text = f"成功加载 {file_path} 的向量"
                    except Exception as e:
                        if spinner:
                            spinner.text = f"加载 {file_path} 向量失败: {str(e)}"
                else:
                    if spinner:
                        spinner.text = f"未找到 {file_path} 的向量缓存"
                        
                    # 从flat_index重建向量
                    if self.flat_index is not None:
                        if spinner:
                            spinner.text = f"从索引重建 {file_path} 的向量..."
                        
                        for doc_idx, chunk_idx in doc_indices:
                            idx = next((i for i, d in enumerate(self.documents) 
                                     if d.metadata['file_path'] == file_path and 
                                     d.metadata['chunk_index'] == chunk_idx), None)
                            
                            if idx is not None:
                                vector = np.zeros((1, self.vector_dim), dtype=np.float32) # type: ignore
                                self.flat_index.reconstruct(idx, vector.ravel())
                                unchanged_vectors.append(vector)

            if not unchanged_vectors:
                if spinner:
                    spinner.text = "未能加载任何未变化文档的向量"
                return None
                
            if spinner:
                spinner.text = f"未变化文档向量加载完成，共 {len(unchanged_vectors)} 个"
                
            return np.vstack(unchanged_vectors)
            
        except Exception as e:
            if spinner:
                spinner.text = f"获取不变向量失败: {str(e)}"
            PrettyOutput.print(f"获取不变向量失败: {str(e)}", OutputType.ERROR)
            return None

    def search(self, query: str, top_k: int = 30) -> List[Tuple[Document, float]]:
        """Search documents with context window"""
        if not self.index:
            self.build_index(self.root_dir)
            
        # 如果索引建立失败或文档列表为空，返回空结果
        if not self.index or len(self.documents) == 0:
            PrettyOutput.print("索引未建立或文档列表为空", OutputType.WARNING)
            return []
            
        # Get query vector
        with yaspin(text="获取查询向量...", color="cyan") as spinner:
            query_vector = get_embedding(self.embedding_model, query)
            query_vector = query_vector.reshape(1, -1)
            spinner.text = "查询向量获取完成"
            spinner.ok("✅")
        
        # Search with more candidates
        with yaspin(text="搜索...", color="cyan") as spinner:
            initial_k = min(top_k * 4, len(self.documents))
            if initial_k == 0:
                spinner.text = "文档为空，搜索终止"
                spinner.fail("❌")
                return []
                
            distances, indices = self.index.search(query_vector, initial_k) # type: ignore
            spinner.text = "搜索完成"
            spinner.ok("✅")
        
        # Process results with context window
        with yaspin(text="处理结果...", color="cyan") as spinner:
            results = []
            seen_files = set()
            
            # 检查索引数组是否为空
            if indices.size == 0 or indices[0].size == 0:
                spinner.text = "搜索结果为空"
                spinner.fail("❌")
                return []
                
            for idx, dist in zip(indices[0], distances[0]):
                if idx != -1 and idx < len(self.documents):  # 确保索引有效
                    doc = self.documents[idx]
                    similarity = 1.0 / (1.0 + float(dist))
                    if similarity > 0.3:
                        file_path = doc.metadata['file_path']
                        if file_path not in seen_files:
                            seen_files.add(file_path)
                            
                            # Get full context from original document
                            original_doc = next((d for d in self.documents 
                                            if d.metadata['file_path'] == file_path), None)
                            if original_doc:
                                window_docs = []  # Add this line to initialize the list
                                full_content = original_doc.content
                                # Find all chunks from this file
                                file_chunks = [d for d in self.documents 
                                            if d.metadata['file_path'] == file_path]
                                # Add all related chunks
                                for chunk_doc in file_chunks:
                                    window_docs.append((chunk_doc, similarity * 0.9))
                            
                            results.extend(window_docs)
                            if len(results) >= top_k * (2 * self.context_window + 1):
                                break
            spinner.text = "处理结果完成"
            spinner.ok("✅")
        
        # Sort by similarity and deduplicate
        with yaspin(text="排序...", color="cyan") as spinner:
            if not results:
                spinner.text = "无有效结果"
                spinner.fail("❌")
                return []
                
            results.sort(key=lambda x: x[1], reverse=True)
            seen = set()
            final_results = []
            for doc, score in results:
                key = (doc.metadata['file_path'], doc.metadata['chunk_index'])
                if key not in seen:
                    seen.add(key)
                    final_results.append((doc, score))
                    if len(final_results) >= top_k:
                        break
            spinner.text = "排序完成"
            spinner.ok("✅")
                    
        return final_results

    def query(self, query: str) -> List[Document]:
        """Query related documents
        
        Args:
            query: Query text
            
        Returns:
            List[Document]: Related documents
        """
        results = self.search(query)
        return [doc for doc, _ in results]

    def ask(self, question: str) -> Optional[str]:
        """Ask questions about documents with enhanced context building"""
        try:
            results = self.search(question)
            if not results:
                return None
            
            prompt = f"""
# 🤖 角色定义
您是一位文档分析专家，能够基于提供的文档提供准确且全面的回答。

# 🎯 核心职责
- 全面分析文档片段
- 准确回答问题
- 引用源文档
- 识别缺失信息
- 保持专业语气

# 📋 回答要求
## 内容质量
- 严格基于提供的文档作答
- 具体且精确
- 在有帮助时引用相关内容
- 指出任何信息缺口
- 使用专业语言

## 回答结构
1. 直接回答
   - 清晰简洁的回应
   - 基于文档证据
   - 专业术语

2. 支持细节
   - 相关文档引用
   - 文件参考
   - 上下文解释

3. 信息缺口（如有）
   - 缺失信息
   - 需要的额外上下文
   - 潜在限制

# 🔍 分析上下文
问题: {question}

相关文档（按相关性排序）：
"""

            # Add context with length control
            with yaspin(text="添加上下文...", color="cyan") as spinner:
                available_count = self.max_token_count - get_context_token_count(prompt) - 1000
                current_count = 0
                
                for doc, score in results:
                    doc_content = f"""
    ## 文档片段 [相关度: {score:.3f}]
    来源: {doc.metadata['file_path']}
    ```
    {doc.content}
    ```
    ---
    """
                    if current_count + get_context_token_count(doc_content) > available_count:
                        PrettyOutput.print(
                            "由于上下文长度限制，部分内容被省略",
                            output_type=OutputType.WARNING
                        )
                        break
                        
                    prompt += doc_content
                    current_count += get_context_token_count(doc_content)

                prompt += """
    # ❗ 重要规则
    1. 仅使用提供的文档
    2. 保持精确和准确
    3. 在相关时引用来源
    4. 指出缺失的信息
    5. 保持专业语气
    6. 使用用户的语言回答
    """
                spinner.text = "添加上下文完成"
                spinner.ok("✅")

            with yaspin(text="正在生成答案...", color="cyan") as spinner:
                model = PlatformRegistry.get_global_platform_registry().get_normal_platform()
                response = model.chat_until_success(prompt)
                spinner.text = "答案生成完成"
                spinner.ok("✅")
                return response
            
        except Exception as e:
            PrettyOutput.print(f"回答失败：{str(e)}", OutputType.ERROR)
            return None

    def is_index_built(self) -> bool:
        """Check if the index is built and valid
        
        Returns:
            bool: True if index is built and valid
        """
        return self.index is not None and len(self.documents) > 0

    def _delete_file_cache(self, file_path: str, spinner=None):
        """Delete cache files for a specific file
        
        Args:
            file_path: Path to the original file
            spinner: Optional spinner for progress information. If None, runs silently.
        """
        try:
            # Delete document cache
            doc_cache_path = self._get_cache_path(file_path, "doc")
            if os.path.exists(doc_cache_path):
                os.remove(doc_cache_path)
                if spinner is not None:
                    spinner.write(f"🗑️ 删除文档缓存: {file_path}")
                    
            # Delete vector cache
            vec_cache_path = self._get_cache_path(file_path, "vec")
            if os.path.exists(vec_cache_path):
                os.remove(vec_cache_path)
                if spinner is not None:
                    spinner.write(f"🗑️ 删除向量缓存: {file_path}")
                    
        except Exception as e:
            if spinner is not None:
                spinner.write(f"❌ 删除缓存失败: {file_path}: {str(e)}")
            PrettyOutput.print(f"删除缓存失败: {file_path}: {str(e)}", output_type=OutputType.ERROR)

def main():
    """Main function"""
    import argparse
    import sys
    
    # Set standard output encoding to UTF-8
    if sys.stdout.encoding != 'utf-8':
        import codecs
        sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
        sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')
    
    parser = argparse.ArgumentParser(description='Document retrieval and analysis tool')
    parser.add_argument('--dir', type=str, help='Directory to process')
    parser.add_argument('--build', action='store_true', help='Build document index')
    parser.add_argument('--search', type=str, help='Search document content')
    parser.add_argument('--ask', type=str, help='Ask about documents')
    args = parser.parse_args()

    try:
        current_dir = os.getcwd()
        rag = RAGTool(current_dir)

        if not args.dir:
            args.dir = current_dir

        if args.dir and args.build:
            rag.build_index(args.dir)
            return 0

        if args.search or args.ask:

            if args.search:
                results = rag.query(args.search)
                if not results:
                    PrettyOutput.print("未找到相关内容", output_type=OutputType.WARNING)
                    return 1
                    
                for doc in results:
                    output = f"""文件: {doc.metadata['file_path']}\n"""
                    output += f"""片段 {doc.metadata['chunk_index'] + 1}/{doc.metadata['total_chunks']}\n"""
                    output += f"""内容:\n{doc.content}\n"""
                    PrettyOutput.print(output, output_type=OutputType.INFO, lang="markdown")
                return 0

            if args.ask:
                # Call ask method
                response = rag.ask(args.ask)
                if not response:
                    PrettyOutput.print("获取答案失败", output_type=OutputType.WARNING)
                    return 1
                    
                # Display answer
                output = f"""{response}"""
                PrettyOutput.print(output, output_type=OutputType.INFO)
                return 0

        PrettyOutput.print("请指定操作参数。使用 -h 查看帮助。", output_type=OutputType.WARNING)
        return 1

    except Exception as e:
        PrettyOutput.print(f"执行失败: {str(e)}", output_type=OutputType.ERROR)
        return 1

if __name__ == "__main__":
    main()
